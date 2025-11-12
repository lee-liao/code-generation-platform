# 在文件顶部添加
from dotenv import load_dotenv
load_dotenv()  # 加载.env文件
import json
import logging
import requests
from fastapi import HTTPException
import os
import oss2
import oss2.exceptions
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.enum.style import WD_STYLE_TYPE 
from docx.enum.text import WD_BREAK
from Markdown2docx import Markdown2docx
from datetime import datetime
from models.report import Report
from sqlmodel import create_engine, Session
import redis
import pypandoc
import copy
from pathlib import Path
BASE_DIR = Path(__file__).parent
REFERENCE_DOCX = str(BASE_DIR / "custom_reference.docx") 

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

OSS_ACCESS_KEY_ID = os.getenv('OSS_ACCESS_KEY_ID')
OSS_ACCESS_KEY_SECRET = os.getenv('OSS_ACCESS_KEY_SECRET')
OSS_ENDPOINT = os.getenv('OSS_ENDPOINT')
OSS_BUCKET_NAME = os.getenv('OSS_BUCKET_NAME')

# Redis配置
REDIS_HOST = os.getenv('REDIS_HOST', 'localhost')
REDIS_PORT = int(os.getenv('REDIS_PORT', 6379))
REDIS_DB = int(os.getenv('REDIS_DB', 0))
REDIS_PASSWORD = os.getenv('REDIS_PASSWORD')

# 创建Redis客户端
redis_client = redis.Redis(
    host=REDIS_HOST,
    port=REDIS_PORT,
    db=REDIS_DB,
    password=REDIS_PASSWORD,
    decode_responses=True
)

DATABASE_URL = os.getenv("DATABASE_URL")
engine = create_engine(
    DATABASE_URL,
    pool_pre_ping=True,  # 添加连接池健康检查
    echo=True  # 开启SQL日志用于调试
)

def upload_to_oss(file_path, object_name=None):
    """上传文件到阿里云OSS"""
    auth = oss2.Auth(OSS_ACCESS_KEY_ID, OSS_ACCESS_KEY_SECRET)
    bucket = oss2.Bucket(auth, OSS_ENDPOINT, OSS_BUCKET_NAME)
    
    if not object_name:
        object_name = os.path.basename(file_path)
    
    try:
        result = bucket.put_object_from_file(object_name, file_path)
        if result.status == 200:
            # 修正URL生成逻辑
            return f"https://file.sflow.pro/{object_name}"
        logging.error(f"OSS上传失败，HTTP状态码: {result.status}")
        return None
    except oss2.exceptions.OssError as e: 
        logging.error(f"OSS服务错误: 状态码[{e.status}] 错误码[{e.code}] 消息[{e.message}]")
    except Exception as e:
        logging.error(f"系统异常: {str(e)}", exc_info=True)
    return None

# Template sections
# template_sections = os.getenv("REPORT_TEMPLATE_SECTIONS").split(",")
template_sections = [
    {
        "title": "公司概况与定位（CPR模型）",
        "prompt": f"""\
请根据以下信息，帮助我提炼企业的商业定位：
1.	我们的核心客户是哪些？
2.	我们的核心产品/服务是什么？
3.	我们目前在产业链中扮演什么角色？（如ODM、解决方案商、平台提供者等）
请用CPR定位模型输出三角结构，说明“我们为谁服务，用什么产品，以什么能力取胜”。
"""
    },
    {
        "title": "宏观趋势与产业全局分析（看全局、判升级）",
        "prompt": f"""\
请基于当前政策、经济、技术和社会趋势，帮助我们总结以下内容：
1.	当前哪些宏观趋势会影响该行业？（PEST分析）
2.	我们所在的行业全景图是怎样的？上下游有哪些关键环节？
3.	是否处于行业升级期？升级的主要表现在哪些方面？
请综合整理成一段战略分析内容，用于公司战略建议书的“趋势洞察”部分。
"""
    },
    {
        "title": "边界扩展分析（破边界）",
        "prompt": f"""\
基于我们现有的客户、产品和能力，请帮助我们识别潜在的增长扩展方向：
1.	我们是否可以用现有产品服务新的客户群？（新客户）
2.	是否可以为当前客户提供新的产品或服务？（新产品）
3.	是否可以用我们的能力承担价值链中新的角色？（新角色）
请以“客户/产品/角色”三维方式输出边界扩展建议。
"""
    },
    {
        "title": "周期识别与战略时机（断周期）",
        "prompt": f"""\
请帮助我们分析当前行业（或我们目标业务）所处的生命周期阶段。
1.	宏观角度：我们所在的行业整体处于什么阶段？（起步期/成长期/成熟期/整合期）
2.	中观角度：我们专注的赛道目前市场集中度和增长速度如何？
3.	微观角度：我们的销售渠道、主要产品、客户是否还在增长周期？
输出一段关于“行业周期识别与战略节奏”的内容。
"""
    },
    {
        "title": "战略机会识别与筛选（择机会）",
        "prompt": f"""\
请帮助我们从以下角度评估战略机会清单，并推荐优先路径：
•	市场吸引力（规模、增长率）
•	公司契合度（资源/能力匹配）
•	投入产出比（资源需求与回报）
•	风险程度（政策、技术、竞争）
请根据这些指标，对机会进行优先级排序，并输出推荐路径。
"""
    },
    {
        "title": "战略定位与目标设定（升定位、立目标）",
        "prompt": f"""\
根据我们选择的战略方向和机会，请帮助我们定义：
1.	新的企业定位（以CPR模型重新定义）
2.	三年/五年期的战略目标，包括：营收目标、利润目标、市场份额、技术或产品目标、客户数量等
请生成一段可用于战略文档的“目标设定与战略定位”章节内容。
"""
    },
    {
        "title": "战略路径与关键举措（定战略、设路径）",
        "prompt": f"""\
请帮我制定一个三年战略路线图，分阶段输出主要举措和资源安排：
1.	第一阶段（第1年）：核心任务和目标
2.	第二阶段（第2年）：主要拓展动作
3.	第三阶段（第3年）：巩固领先地位的策略
同时，请说明每一阶段的重点资源配置（技术、人才、渠道、资金）。
"""
    },
    {
        "title": "可视化图表与知识图谱（可选）",
        "prompt": f"""\
基于我们完整的战略分析内容，请帮助我们设计对应的图表和结构图，包括：
•	产业全景图
•	战略目标蓝图
•	边界扩展三轴图
•	赛道周期判断曲线图
•	战略路线图（Roadmap）
并提炼出一套知识图谱结构（企业 → 赛道、客户 → 需求、能力 → 角色等）以供AI模型调用。
"""
    }
]

# Function to query Perplexity API for company information
def generate_report_section_by_perplexity(prompt, section_name):
    try:
        logging.info(f"Generating report section by perplexity: {section_name}")
        api_key = os.getenv('PERPLEXITY_KEY')
        url = "https://api.perplexity.ai/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "accept": "application/json",
            "content-type": "application/json"
        }
        payload = {
            "model": "sonar-pro",
            "messages": [
                {"role": "system", "content": f"""\
                  你是一位专业商业报告撰写助手，你用用户提供的公司名称，还有网页信息搜索互联网，按照海峰老师的模版，提供一个公司战略分析师的相关建议，你可以参考网络上战略分析师的一些信息，还有企业管理，发展战略的一些专业知识，给用户提供最有价值的商务建议，有的建议，可能不在用户的回答里面，作为战略分析师，你采用你的专业高度，尽量给用户提供对企业发展有用的建议，可以通过搜索互联网找到类似的战略分析建议，来帮助公司发展战略。请遵守以下格式要
                             1. 列表项使用Markdown的无序列表语法（以 - 或 * 开头）
                             2. 避免使用自动编号
                             3. 章节标题使用##符号
                             4. 关键术语使用**加粗**强调
                             5. 禁止使用数字编号列表
请参考一亿中流网站关于海峰老师的战略“https://www.yyzltop.com/”，还有其他重要的企业发展战略图书报告中文和英文的书籍或者网站
                             """
                             },
                {"role": "user", "content": prompt}
            ],
            'max_tokens': 2000,
            'temperature': 0.7
        }
        response = requests.post(url, json=payload, headers=headers)
        logging.info(f"Used tokens: {response.json().get('usage', {}).get('total_tokens', 0)}")
        if response.status_code == 200:
            content = (response.json().get('choices', [{}])[0]
                      .get('message', {})
                      .get('content', '').strip())
            logging.info(f"Successfully generated section by perplexity: {section_name}")
            logging.info(f"Successfully generated section by perplexity content: {content}")
            return content
        else:
            logging.error(f"Failed to retrieve data from Perplexity API. Status code: {response.status_code}. Response: {response.text}")
            raise HTTPException(status_code=500, detail=f"Failed to retrieve data from Perplexity API. Status code: {response.status_code}")
    except Exception as e:
        logging.error(f"Error querying Perplexity API: {str(e)}")
        logging.error(f"Error generating report section '{section_name}': {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="Error querying Perplexity API")

def get_report_url(session_id):
    """从Redis获取指定session_id的OSS URL"""
    try:
        redis_key = f"report:hf:{session_id}"
        oss_url = redis_client.get(redis_key)
        if oss_url:
            logging.info(f"成功从Redis获取OSS URL，session_id: {session_id}")
            return {
                "message": "Get report successfully",
                "session_id": session_id,
                "file_url": oss_url
            }
        logging.warning(f"未找到对应的OSS URL，session_id: {session_id}")
        return {
            "message": "Currently no report available",
            "session_id": session_id,
            "file_url": ""
        }
    except Exception as e:
        logging.error(f"Redis查询失败: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="暂时无法获取报告链接，请稍后重试"
        )


def generate_business_report(file_url, session_id, user_id):
    logging.info("Starting business report generation")
    
    try:
        # 下载数据部分保持不变
        logging.info(f"Downloading report data from: {file_url}")
        import requests
        response = requests.get(file_url)
        response.raise_for_status()
        data = response.json()
        logging.info("Report data downloaded successfully")
        try:
            redis_key1 = f"report:qa:{session_id}"
            redis_key2 = f"report:type:{session_id}"
            redis_client.set(redis_key1, json.dumps(data))
            redis_client.set(redis_key2, "haifeng")
            logging.info(f"成功存储QA到Redis，键名: {redis_key1}")
        except Exception as e:
            logging.error(f"Redis存储失败: {str(e)}", exc_info=True)
        company_name = data.get("您的企业名称是什么？", "")
        website_url = data.get("您的网站网址是什么？", "")
        logging.info(f"Company Name: {company_name}, Website URL: {website_url}")

        # 创建新文档
        document = Document()

        # ========== 全局样式设置 ==========
        normal_style = document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'SimSun'
        normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        normal_font.size = Pt(12)

        # 新增列表样式配置
        list_style = document.styles.add_style('BulletList', WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = document.styles['List Bullet']
        list_font = list_style.font
        list_font.name = 'SimSun'
        list_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        list_font.size = Pt(12)
        
        # ========== 添加报告标题 ==========
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(f"{company_name}企业战略分析与建议")
        title_font = title_run.font
        title_font.name = 'SimHei'
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        title_font.size = Pt(16)
        title_font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(24)
        who = f"我们的公司名称是：{company_name}, 网站地址是：{website_url}"
        
        # ========== 生成报告正文 ==========
        for section_idx, section in enumerate(template_sections):
            # 添加章节标题（黑体14pt）
            heading_paragraph = document.add_paragraph()
            heading_run = heading_paragraph.add_run(f"第{section_idx+1}章 {section['title']}")
            heading_font = heading_run.font
            heading_font.name = 'SimHei'
            heading_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_paragraph.paragraph_format.space_before = Pt(12)
            heading_paragraph.paragraph_format.space_after = Pt(12)

            # 获取Markdown内容
            markdown_content = generate_report_section_by_perplexity(f"{who}, {section['prompt']}", section['title'])
            
            # 使用pypandoc转换
            temp_docx = f"temp_{session_id}_{section_idx}.docx"
            try:
                # 将Markdown转换为临时docx文件
                pypandoc.convert_text(
                    markdown_content,
                    'docx',
                    format='md',
                    outputfile=temp_docx,
                    encoding='utf-8',
                    extra_args=['--reference-doc', REFERENCE_DOCX]
                )
                
                # 合并转换后的内容到当前章节标题之后
                temp_doc = Document(temp_docx)
                current_parent = heading_paragraph._element.getparent()
                current_element = heading_paragraph._element
                
                for element in temp_doc.element.body:
                    new_element = copy.deepcopy(element)
                    current_element.addnext(new_element)
                    current_element = new_element

            finally:
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)

            # 保存原始Markdown到数据库
            new_report = Report(
                user_id=user_id,
                session_id=session_id,
                title=section['title'],
                subtitle="",
                content=markdown_content
            )

            # 数据库操作
            with Session(engine) as session:
                try:
                    session.begin()
                    session.add(new_report)
                    session.commit()
                    session.refresh(new_report)
                except Exception as e:
                    session.rollback()
                    logging.error(f"Database error: {str(e)}")
                    raise

            # 添加分页符（在章节内容末尾）
            if section_idx != len(template_sections)-1:
                document.add_page_break()

        # ========== 保存与上传 ==========
        local_filename = f"{session_id}_haifeng_business_report.docx"
        document.save(local_filename)
        
        # OSS上传
        object_name = f"reports/{datetime.now().strftime('%Y%m%d')}/{local_filename}"
        oss_url = upload_to_oss(local_filename, object_name)
        
        if not oss_url:
            raise Exception("文件上传到OSS失败")
        
        # 清理本地文件
        if os.path.exists(local_filename):
            os.remove(local_filename)
        logging.info(f"文件已成功上传至OSS，链接是: {oss_url}")
        
        # 存储到Redis，设置24小时过期时间
        try:
            redis_key = f"report:hf:{session_id}"
            redis_client.set(redis_key, oss_url)
            logging.info(f"成功存储OSS URL到Redis，键名: {redis_key}")
        except Exception as e:
            logging.error(f"Redis存储失败: {str(e)}", exc_info=True)
        
        logging.info("Successfully generated business report document")
        return {
            "message": "Business report generated successfully",
            "session_id": session_id,
            "file_url": oss_url
        }
        
    except Exception as e:
        logging.error(f"Error generating business report: {str(e)}", exc_info=True)
        return {
            "status": "error",
            "message": str(e)
        }


def old1generate_business_report(file_url, session_id, user_id):
    logging.info("Starting business report generation")
    
    try:
        # 下载数据部分保持不变
        logging.info(f"Downloading report data from: {file_url}")
        import requests
        response = requests.get(file_url)
        response.raise_for_status()
        data = response.json()
        logging.info("Report data downloaded successfully")
        company_name = data.get("您的企业名称是什么？", "")
        website_url = data.get("您的网站网址是什么？", "")
        logging.info(f"Company Name: {company_name}, Website URL: {website_url}")

        # 创建新文档
        document = Document()

        # ========== 全局样式设置 ==========
        normal_style = document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'SimSun'
        normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        normal_font.size = Pt(12)
        
        # ========== 添加报告标题 ==========
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(f"{company_name}企业战略分析与建议")
        title_font = title_run.font
        title_font.name = 'SimHei'
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        title_font.size = Pt(16)
        title_font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(24)
        who = f"我们的公司名称是：{company_name}, 网站地址是：{website_url}"
        
        # ========== 生成报告正文 ==========
        for section_idx, section in enumerate(template_sections):
            # 添加章节标题（黑体14pt）
            heading_paragraph = document.add_paragraph()
            heading_run = heading_paragraph.add_run(f"第{section_idx+1}章 {section['title']}")
            heading_font = heading_run.font
            heading_font.name = 'SimHei'
            heading_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_paragraph.paragraph_format.space_before = Pt(12)
            heading_paragraph.paragraph_format.space_after = Pt(12)

            # 获取Markdown内容
            markdown_content = generate_report_section_by_perplexity(f"{who}, {section['prompt']}", section['title'])
            
            # 使用pypandoc转换
            temp_docx = f"temp_{session_id}_{section_idx}.docx"
            try:
                # 将Markdown转换为临时docx文件
                pypandoc.convert_text(
                    markdown_content,
                    'docx',
                    format='md',
                    outputfile=temp_docx,
                    encoding='utf-8'
                )
                
                # 合并转换后的内容到主文档
                temp_doc = Document(temp_docx)
                for element in temp_doc.element.body:
                    document.element.body.append(element)
                
            finally:
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)

            # 保存原始Markdown到数据库
            new_report = Report(
                user_id=user_id,
                session_id=session_id,
                title=section['title'],
                subtitle="",
                content=markdown_content
            )

            # 数据库操作
            with Session(engine) as session:
                try:
                    session.begin()
                    session.add(new_report)
                    session.commit()
                    session.refresh(new_report)
                except Exception as e:
                    session.rollback()
                    logging.error(f"Database error: {str(e)}")
                    raise

            # 添加分页符（在转换后的内容末尾）
            if section_idx != len(template_sections)-1:
                if document.paragraphs:
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.add_run().add_break(WD_BREAK.PAGE)

        # ========== 保存与上传 ==========
        local_filename = f"{session_id}_haifeng_business_report.docx"
        document.save(local_filename)
        
        # OSS上传
        object_name = f"reports/{datetime.now().strftime('%Y%m%d')}/{local_filename}"
        oss_url = upload_to_oss(local_filename, object_name)
        
        if not oss_url:
            raise Exception("文件上传到OSS失败")
        
        # 清理本地文件
        if os.path.exists(local_filename):
            os.remove(local_filename)
        logging.info(f"文件已成功上传至OSS，链接是: {oss_url}")
        # 存储到Redis，设置24小时过期时间
        try:
            redis_key = f"report:hf:{session_id}"
            redis_client.set(redis_key, oss_url)
            logging.info(f"成功存储OSS URL到Redis，键名: {redis_key}")
        except Exception as e:
            logging.error(f"Redis存储失败: {str(e)}", exc_info=True)
        
        logging.info("Successfully generated business report document")
        return {
            "message": "Business report generated successfully",
            "session_id": session_id,
            "file_url": oss_url
        }
        
    except Exception as e:
        logging.error(f"Error generating business report: {str(e)}", exc_info=True)
        return {
            "status": "error",
            "message": str(e)
        }
