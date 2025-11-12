# 在文件顶部添加
from dotenv import load_dotenv
load_dotenv()  # 加载.env文件
import json
import openai
import logging
import os
import oss2
import oss2.exceptions
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.enum.text import WD_BREAK
from datetime import datetime
from models.report import Report
from sqlmodel import create_engine, Session

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# OpenAI API setup
client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# 在OpenAI API setup之后添加OSS配置
OSS_ACCESS_KEY_ID = os.getenv('OSS_ACCESS_KEY_ID')
OSS_ACCESS_KEY_SECRET = os.getenv('OSS_ACCESS_KEY_SECRET')
OSS_ENDPOINT = os.getenv('OSS_ENDPOINT')
OSS_BUCKET_NAME = os.getenv('OSS_BUCKET_NAME')

DATABASE_URL = os.getenv("DATABASE_URL")
engine = create_engine(
    DATABASE_URL,
    pool_pre_ping=True,  # 添加连接池健康检查
    echo=True  # 开启SQL日志用于调试
)

def upload_to_oss(file_path, object_name=None):
    """上传文件到阿里云OSS"""
    auth = oss2.Auth(OSS_ACCESS_KEY_ID, OSS_ACCESS_KEY_SECRET)
    bucket = oss2.Bucket(auth, OSS_ENDPOINT, OSS_BUCKET_NAME)
    
    if not object_name:
        object_name = os.path.basename(file_path)
    
    try:
        result = bucket.put_object_from_file(object_name, file_path)
        if result.status == 200:
            # 修正URL生成逻辑
            return f"https://file.sflow.pro/{object_name}"
        logging.error(f"OSS上传失败，HTTP状态码: {result.status}")
        return None
    except oss2.exceptions.OssError as e: 
        logging.error(f"OSS服务错误: 状态码[{e.status}] 错误码[{e.code}] 消息[{e.message}]")
    except Exception as e:
        logging.error(f"系统异常: {str(e)}", exc_info=True)
    return None

# Template sections
# template_sections = os.getenv("REPORT_TEMPLATE_SECTIONS").split(",")
template_sections = [
    {
        "title": "执行摘要",
        "subtitles": [
            "业务概述",
            "主要亮点与成就",
            "报告目标",
            "关键发现与建议摘要"
        ]
    },
    {
        "title": "公司概述",
        "subtitles": [
            "公司名称与描述",
            "使命、愿景和价值观",
            "商业模式与收入来源",
            "法律结构与所有权",
            "组织架构与关键人员"
        ]
    },
    {
        "title": "行业与市场分析",
        "subtitles": [
            "行业概述与趋势",
            "市场规模与增长潜力",
            "目标市场与客户细分",
            "竞争格局与主要竞争对手",
            "影响企业的法规与经济因素"
        ]
    },
    {
        "title": "业务运营",
        "subtitles": [
            "产品/服务提供",
            "供应链与物流",
            "技术与基础设施",
            "关键合作伙伴与联盟",
            "运营挑战与解决方案"
        ]
    },
    {
        "title": "营销与销售策略",
        "subtitles": [
            "市场定位与品牌策略",
            "客户获取渠道",
            "定价策略",
            "促销与广告策略",
            "销售渠道与客户留存策略"
        ]
    },
    {
        "title": "财务概述",
        "subtitles": [
            "收入来源与定价模式",
            "损益概览",
            "现金流分析",
            "资金需求与投资策略",
            "财务预测与增长预期"
        ]
    },
    {
        "title": "商业风险与挑战",
        "subtitles": [
            "SWOT 分析（优势、劣势、机会、威胁）",
            "行业特定风险",
            "经济与政治风险",
            "应急计划与风险缓解策略"
        ]
    },
    {
        "title": "业务增长与扩展计划",
        "subtitles": [
            "短期与长期增长战略",
            "市场扩展与多元化发展",
            "创新与研发（R&D）",
            "战略合作伙伴关系与并购",
            "国际扩展计划（如适用）"
        ]
    },
    {
        "title": "可持续发展与企业社会责任（CSR）",
        "subtitles": [
            "环境影响与可持续发展措施",
            "伦理商业实践",
            "社区参与与社会责任",
            "公司治理与合规"
        ]
    },
    {
        "title": "结论与建议",
        "subtitles": [
            "主要发现总结",
            "关键见解与商业洞察",
            "可执行建议",
            "下一步计划与实施方案"
        ]
    },
    {
        "title": "附录与支持文档",
        "subtitles": [
            "财务报表与预测",
            "市场研究数据",
            "法律文件与许可证",
            "组织架构图",
            "其他相关报告或调查"
        ]
    }
]


# Function to generate report section using OpenAI API
def generate_report_section(section_name, data):
    try:
        logging.info(f"Generating report section: {section_name}")
        prompt = f"Generate a paragraph for the '{section_name}' section of a business report based on the following data: {data}"
        response = client.chat.completions.create(
            # model="gpt-4-turbo",
            model="gpt-4o-mini",
            #model="gpt-4o",
            messages=[{"role": "system", "content": "You are an expert business report writer."},
                      {"role": "user", "content": prompt}],
            max_tokens=1024,
            temperature=0.7
        )
        
        result = response.choices[0].message.content.strip()
        logging.info(f"Successfully generated section: {section_name}")
        return result
    
    except Exception as e:
        logging.error(f"Error generating report section '{section_name}': {str(e)}", exc_info=True)
        return "Error generating section. Please check logs."

# Function to generate the business report document
def generate_business_report(file_url, session_id, user_id):
    logging.info("Starting business report generation")
    
    try:
        # 下载数据部分保持不变
        logging.info(f"Downloading report data from: {file_url}")
        import requests
        response = requests.get(file_url)
        response.raise_for_status()
        data = response.json()
        logging.info("Report data downloaded successfully")
        
        # 创建新文档
        document = Document()
        
        # ========== 全局样式设置 ==========
        normal_style = document.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'SimSun'
        normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        normal_font.size = Pt(12)
        
        # ========== 添加报告标题 ==========
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run("业务分析报告")
        title_font = title_run.font
        title_font.name = 'SimHei'
        title_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        title_font.size = Pt(16)
        title_font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(24)
        
        # ========== 生成报告正文 ==========
        for section_idx, section in enumerate(template_sections):
            # 添加章节标题（黑体14pt）
            heading_paragraph = document.add_paragraph()
            heading_run = heading_paragraph.add_run(f"第{section_idx+1}章 {section['title']}")
            heading_font = heading_run.font
            heading_font.name = 'SimHei'
            heading_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_paragraph.paragraph_format.space_before = Pt(12)
            heading_paragraph.paragraph_format.space_after = Pt(12)
            
            for subtitle_idx, subtitle in enumerate(section['subtitles']):
                # 添加子标题（宋体12pt，带项目符号）
                content_paragraph = document.add_paragraph()
                run = content_paragraph.add_run(f"• {subtitle}")
                content_font = run.font
                content_font.name = 'SimSun'
                content_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                content_font.size = Pt(12)
                
                # 添加子标题内容
                content_paragraph = document.add_paragraph()
                run = content_paragraph.add_run(generate_report_section(subtitle, data))
                content_font = run.font
                content_font.name = 'SimSun'
                content_font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                content_font.size = Pt(12)
                content_paragraph.paragraph_format.line_spacing = 1.5
                content_paragraph.paragraph_format.space_after = Pt(18)
                logging.info(f"session_id: {session_id}, title: {section['title']}, subtitle: {section['subtitles'][subtitle_idx]}, content: {content_paragraph.text}")
                new_report = Report(
                    user_id=user_id,
                    session_id=session_id,
                    title=section['title'],
                    subtitle=section['subtitles'][subtitle_idx],
                    content=content_paragraph.text
                )
                # 使用数据库会话添加数据
               # 修改后的数据库操作部分
                with Session(engine) as session:
                    try:
                        session.begin()  # 显式开始事务
                        session.add(new_report)
                        session.commit()
                        session.refresh(new_report)
                    except Exception as e:
                        session.rollback()
                        logging.error(f"Database error: {str(e)}")
                        raise
                # 在最后一个子标题的内容段落添加分页符
                if section_idx != len(template_sections)-1 and subtitle_idx == len(section['subtitles'])-1:
                    content_paragraph.add_run().add_break(WD_BREAK.PAGE)

        # ========== 保存与上传 ==========
        local_filename = f"{session_id}_business_report.docx"
        document.save(local_filename)
        # ... 后续上传和返回逻辑保持不变 ...
        
        # OSS上传逻辑保持不变
        object_name = f"reports/{datetime.now().strftime('%Y%m%d')}/{local_filename}"
        oss_url = upload_to_oss(local_filename, object_name)
        
        if not oss_url:
            raise Exception("文件上传到OSS失败")
        
        # 清理本地文件
        if os.path.exists(local_filename):
            os.remove(local_filename)
        
        logging.info(f"文件已成功上传至OSS，链接是: {oss_url}")
        logging.info("Successfully generated business report document")
        
        return {
            "message": "Business report generated successfully",
            "session_id": session_id,
            "file_url": oss_url
        }
        
    except Exception as e:
        logging.error(f"Error generating business report: {str(e)}", exc_info=True)
        return {
            "status": "error",
            "message": str(e)
        }